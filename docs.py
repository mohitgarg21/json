from docx import Document
from docx.shared import Pt


document = Document()

p = document.add_heading('\t\t\t\tTerms and Conditions', 4)
p.add_run().arial = True

p = document.add_paragraph(' Your grade/band and designation will be as per details mentioned i this letter. However, it may be noted that '
                           'designations jobs grades/bands could change  from time depending on Company requirements and in '
                           'accordance with HR policies. ', style='List Number')
p.add_run().arial = True
p.paragraph_format.space_after = Pt(30)

p = document.add_paragraph(' The Company may transfer you across function / location / affiliates as per its business requirements. You may be '
                           'required to work on matters pertaining to the Company and /or its affitiates, as decided and communicated by the '
                           'Company from time to time.', style='List Number')

p.add_run().arial = True

document.add_page_break()

document.save('pythonDocs.docx')